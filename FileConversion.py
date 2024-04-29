import tkinter as tk
from tkinter import filedialog, messagebox
import os
import spacy
import matplotlib.pyplot as plt
from docx import Document
from docx.shared import Pt
from docx.enum.text import WD_PARAGRAPH_ALIGNMENT
from docx.oxml import OxmlElement
from bs4 import BeautifulSoup
import PyPDF2


#BEGIN PDF CODE
# Load the spaCy English model
nlp = spacy.load("en_core_web_sm")

def extract_headers(text):
    doc = nlp(text)
    headers = []
    for sent in doc.sents:
        # Check if the sentence starts with a capital letter and has a colon
        if sent.text.strip() and sent.text[0].isupper() and ':' in sent.text:
            headers.append(sent.text.strip())
        # Or check if the sentence starts with a capital letter and has a length longer than a certain threshold
        elif sent.text.strip() and sent.text[0].isupper() and len(sent.text.split()) > 3:
            headers.append(sent.text.strip())
    return "\n".join(headers)


def extract_function(text):
    doc = nlp(text)
    # Find mathematical expressions using dependency parsing
    math_expressions = [token.text for token in doc if token.dep_ == "amod"]
    return " ".join(math_expressions)


# Function to open a file dialog and get the file path for PDF conversion
def open_file_for_pdf():
    file_path = filedialog.askopenfilename(filetypes=[("Text files", "*.txt")])
    if file_path:
        convert_to_pdf(file_path)
        messagebox.showinfo("Success", f"PDF file has been created successfully!")

def convert_to_pdf(text_file):
    # Read the text file
    with open(text_file, 'r') as file:
        function_text = file.read()

    # Extract headers
    headers = extract_headers(function_text)

    # Extract the function
    function = extract_function(function_text)

    # Create a plot using matplotlib with a fixed figure size
    fig, ax = plt.subplots(figsize=(8.27, 11.69))  # A4 size (approx.)

    # Check if there are headers
    if headers:
        # Render headers with a larger font size and left alignment
        ax.text(0.1, 0.9, headers, fontsize=16, ha='left')
    # Check if the extracted function is empty (no mathematical expressions)
    elif function:
        # Use LaTeX rendering for mathematical expressions
        ax.text(0.1, 0.5, f"${function}$", fontsize=12, ha='left')
    else:
        # Use default text rendering for general text
        ax.text(0.1, 0.5, function_text, fontsize=12, ha='left')

    ax.axis('off')

    # Save the plot as a PDF file
    pdf_file = os.path.splitext(text_file)[0] + '.pdf'
    plt.savefig(pdf_file, bbox_inches='tight', pad_inches=0.1)
    plt.close()

    print(f"PDF file saved as: {pdf_file}")
#END PDF CODE

#BEGIN DOCX CODE
# Function to open a file dialog and get the file path for Word conversion
def open_file_for_docx():
    file_path = filedialog.askopenfilename(filetypes=[("Text files", "*.txt")])
    if file_path:
        convert_to_docx(file_path)
        messagebox.showinfo("Success", f"Word document has been created successfully!")

def apply_heading_style(paragraph):
    paragraph.style = 'Heading 1'

def apply_list_style(paragraph):
    paragraph.style = 'List Bullet'

def apply_title_style(paragraph):

  paragraph.alignment = WD_PARAGRAPH_ALIGNMENT.CENTER

  paragraph_format = paragraph.paragraph_format
  paragraph_format.space_before = Pt(24)

  run = paragraph.runs[0]
  font = run.font
  font.bold = True
  font.size = Pt(24)



def convert_to_docx(text_file):

  document = Document()

  # Read first line and add as title
  with open(text_file) as file:
    first_line = next(file).strip()
    paragraph = document.add_paragraph(first_line)
    apply_title_style(paragraph)

  # Read rest of lines    
  with open(text_file) as file:  
    for line in file:
      line = line.strip()

      if not line:
        continue 

      if line.startswith('#'):
        paragraph = document.add_paragraph(line.strip('#').strip())
        apply_heading_style(paragraph)

      elif line.startswith('- '):  
        paragraph = document.add_paragraph(line.strip('- ').strip())
        apply_list_style(paragraph)

      else:
        document.add_paragraph(line)

      document.add_paragraph()

  # Save docx file
  docx_file = os.path.splitext(text_file)[0] + '.docx'
  document.save(docx_file)
  print(f"Word document saved as: {docx_file}")




#BEGIN HTML CODE
# Function to open a file dialog and get the file path for HTML conversion
def open_file_for_html():
    file_path = filedialog.askopenfilename(filetypes=[("Text files", "*.txt")])
    if file_path:
        convert_to_html(file_path)
        messagebox.showinfo("Success", f"HTML file has been created successfully!")

def convert_to_html(text_file):
    with open(text_file, 'r') as file:
        content = file.read()
    
    # Escape HTML characters in content
    content = content.replace('&', '&amp;').replace('<', '&lt;').replace('>', '&gt;')

    # Wrap content in basic HTML structure
    html_content = f'''<!DOCTYPE html>
<html lang="en">
<head>
    <meta charset="UTF-8">
    <meta name="viewport" content="width=device-width, initial-scale=1.0">
    <title>Document</title>
</head>
<body>
<pre>{content}</pre>
</body>
</html>
'''

    html_file = os.path.splitext(text_file)[0] + '.html'
    with open(html_file, 'w') as file:
        file.write(html_content)
    
    print(f"HTML file saved as: {html_file}")
#END HTML CODE

#BEGIN APP CODE
# Set up the main application window
root = tk.Tk()
root.title("Text-File Converter")

# Window set-up
root.geometry('600x300')  # Width x Height
root.configure(bg='White')

# Configure the button size and font size to make them bigger
button_options = {
    'padx': 10,  # Add some padding on x axis
    'pady': 10,  # Add some padding on y axis
    'font': ('Roboto', 16),  # Increase the font size
    'bg': 'White'
}

# HTML to text
def open_file_for_html_to_text():
    file_path = filedialog.askopenfilename(filetypes=[("HTML files", "*.html")])
    if file_path:
        text = convert_html_to_text(file_path)
        messagebox.showinfo("HTML to Text", f"Extracted Text:\n\n{text[:500]}")

def convert_html_to_text(html_file):
    with open(html_file, 'r') as file:
        soup = BeautifulSoup(file, 'html.parser')
    text = soup.get_text()
    txt_file = os.path.splitext(html_file)[0] + '_text.txt'
    with open(txt_file, 'w') as file:
        file.write(text)
    print(f"Text file saved as: {txt_file}")
    return text
# end

# PDF to text
def open_file_for_pdf_to_text():
    file_path = filedialog.askopenfilename(filetypes=[("PDF files", "*.pdf")])
    if file_path:
        text = convert_pdf_to_text(file_path)
        messagebox.showinfo("PDF to Text", f"Extracted Text:\n\n{text[:500]}")

def convert_pdf_to_text(pdf_file):
    with open(pdf_file, 'rb') as file:
        reader = PyPDF2.PdfReader(file)
        text = ''
        for page_num in range(len(reader.pages)): 
            page = reader.pages[page_num]
            text += page.extract_text()
    txt_file = os.path.splitext(pdf_file)[0] + '_text.txt'
    with open(txt_file, 'w') as file:
        file.write(text)
    print(f"Text file saved as: {txt_file}")
    return text
# end

# word to text
def open_file_for_docx_to_text():
    file_path = filedialog.askopenfilename(filetypes=[("Word files", "*.docx")])
    if file_path:
        text = convert_docx_to_text(file_path)
        messagebox.showinfo("DOCX to Text", f"Extracted Text:\n\n{text[:500]}")

def convert_docx_to_text(docx_file):
    doc = Document(docx_file)
    text = ''
    for para in doc.paragraphs:
        text += para.text + '\n'
    txt_file = os.path.splitext(docx_file)[0] + '_text.txt'
    with open(txt_file, 'w') as file:
        file.write(text)
    print(f"Text file saved as: {txt_file}")
    return text
# end

# Additional Buttons in the GUI for the new functionalities
open_button_html_to_text = tk.Button(root, text="Convert HTML to Text", command=open_file_for_html_to_text, **button_options)
open_button_html_to_text.pack(expand=True)

open_button_pdf_to_text = tk.Button(root, text="Convert PDF to Text", command=open_file_for_pdf_to_text, **button_options)
open_button_pdf_to_text.pack(expand=True)

# Additional button for DOCX to Text functionality
open_button_docx_to_text = tk.Button(root, text="Convert Word to Text", command=open_file_for_docx_to_text, **button_options)
open_button_docx_to_text.pack(expand=True)



# Add a button to open the file dialog for PDF conversion
open_button_pdf = tk.Button(root, text="Convert Text to PDF", command=open_file_for_pdf, **button_options)
open_button_pdf.pack(expand=True)

# Add a button to open the file dialog for Word conversion
open_button_docx = tk.Button(root, text="Convert Text to Word", command=open_file_for_docx, **button_options)
open_button_docx.pack(expand=True)

# Add a button for html conversion.
open_button_html = tk.Button(root, text="Convert Text to HTML", command=open_file_for_html, **button_options)
open_button_html.pack(expand=True)

# Start the Tkinter event loop
root.mainloop()
#END APP CODE